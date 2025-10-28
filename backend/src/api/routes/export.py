from __future__ import annotations

import io
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, Optional
import re

import markdown as mdlib
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel, Field

from src.services.markdowner import blocks_to_markdown

# 添加可选依赖检查
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
    print("成功加载reportlab库，备用PDF导出功能已启用")
except ImportError:
    HAS_REPORTLAB = False
    print("提示: 安装reportlab库可以提供备用PDF导出功能: pip install reportlab")

# 检查是否有中文字体可用
def _register_cjk_fonts():
    """注册中文字体供reportlab使用"""
    if not HAS_REPORTLAB:
        return False
        
    # 尝试常见的中文字体路径
    font_paths = [
        # Windows中文字体
        "C:/Windows/Fonts/simsun.ttc",  # 宋体
        "C:/Windows/Fonts/simhei.ttf",   # 黑体
        "C:/Windows/Fonts/simkai.ttf",   # 楷体
        # Linux中文字体
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/wqy-microhei/wqy-microhei.ttc",
        # macOS中文字体
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf"
    ]
    
    for font_path in font_paths:
        if os.path.exists(font_path):
            try:
                if font_path.endswith("simsun.ttc"):
                    pdfmetrics.registerFont(TTFont("SimSun", font_path, subfontIndex=0))
                    return "SimSun"
                else:
                    font_name = os.path.basename(font_path).split('.')[0]
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    return font_name
            except Exception as e:
                print(f"注册字体失败: {font_path}, 错误: {e}")
    
    # 没有找到任何可用字体
    return False

router = APIRouter(prefix="/export", tags=["export"])

# 修改默认字体为宋体
DEFAULT_CJK_FONT = "宋体"
DEFAULT_FONT_SIZE_PT = 12.0
DEFAULT_HEADING_COLOR = "000000"
DEFAULT_QUOTE_BORDER_COLOR = "D1D5DB"
MARKDOWN_CSS_NAME = "markdown.css"


class DocxStyleOptions(BaseModel):
    font_name: Optional[str] = Field(None, description="段落主字体名称")
    font_size_pt: Optional[float] = Field(None, gt=4, lt=48, description="段落字体大小（pt）")
    heading_font_name: Optional[str] = Field(None, description="标题字体名称")
    heading_color: Optional[str] = Field(None, description="标题颜色，Hex 形式")
    line_spacing: Optional[float] = Field(None, ge=1.0, le=3.0, description="段落行距倍数")
    paragraph_spacing_before: Optional[float] = Field(None, ge=0, le=60, description="段前间距（pt）")
    paragraph_spacing_after: Optional[float] = Field(None, ge=0, le=60, description="段后间距（pt）")
    quote_border_color: Optional[str] = Field(None, description="引用块左侧竖线颜色 Hex")


class ExportRequest(BaseModel):
    blocks: list[Dict[str, Any]]
    title: Optional[str] = None
    docx_options: Optional[DocxStyleOptions] = None


def _sanitize_hex_color(candidate: Optional[str], fallback: str) -> str:
    if not candidate:
        return fallback
    value = candidate.strip().lstrip("#")
    if len(value) == 3:
        value = "".join(ch * 2 for ch in value)
    if len(value) != 6:
        return fallback
    try:
        int(value, 16)
    except ValueError:
        return fallback
    return value.upper()


def _apply_normal_style(document, opts: Dict[str, Any]) -> None:
    try:
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore
        from docx.shared import RGBColor  # type: ignore
    except Exception:
        return

    font_name = opts.get("font_name", DEFAULT_CJK_FONT)
    font_size_pt = opts.get("font_size_pt", DEFAULT_FONT_SIZE_PT)

    try:
        normal = document.styles["Normal"]
    except Exception:
        return

    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(font_size_pt)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    paragraph_format = getattr(normal, "paragraph_format", None)
    if paragraph_format:
        if opts.get("line_spacing") is not None:
            paragraph_format.line_spacing = opts["line_spacing"]
        if opts.get("paragraph_spacing_before") is not None:
            paragraph_format.space_before = Pt(opts["paragraph_spacing_before"])
        if opts.get("paragraph_spacing_after") is not None:
            paragraph_format.space_after = Pt(opts["paragraph_spacing_after"])


def _apply_heading_styles(document, opts: Dict[str, Any]) -> None:
    try:
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import RGBColor  # type: ignore
    except Exception:
        return

    heading_font = opts.get("heading_font_name") or opts.get("font_name") or DEFAULT_CJK_FONT
    heading_color = _sanitize_hex_color(opts.get("heading_color"), DEFAULT_HEADING_COLOR)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[style_name]
            style.font.name = heading_font
            style._element.rPr.rFonts.set(qn("w:eastAsia"), heading_font)
            
            # 标题和一级标题使用等线字体
            if style_name in ("Title", "Heading 1"):  
                style.font.name = "等线"
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            # 说话人部分通常是 Heading 3，使用宋体
            elif style_name in ("Heading 3"):  
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            try:
                style.font.color.rgb = RGBColor.from_string(heading_color)
            except Exception:
                pass
        except Exception:
            continue


def _set_quote_border(paragraph, color_hex: str) -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        return
    p_elem = paragraph._p
    props = p_elem.pPr
    if props is None:
        props = OxmlElement("w:pPr")
        p_elem.append(props)
    border = props.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        props.append(border)
    left = border.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        border.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color_hex)


def _apply_paragraph_format(paragraph, opts: Dict[str, Any], *, is_quote: bool = False) -> None:
    try:
        from docx.shared import Pt  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception:
        return

    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if not paragraph_format:
        return
    if opts.get("line_spacing") is not None:
        paragraph_format.line_spacing = opts["line_spacing"]
    if opts.get("paragraph_spacing_before") is not None:
        paragraph_format.space_before = Pt(opts["paragraph_spacing_before"])
    if opts.get("paragraph_spacing_after") is not None:
        paragraph_format.space_after = Pt(opts["paragraph_spacing_after"])
    if is_quote:
        color = _sanitize_hex_color(opts.get("quote_border_color"), DEFAULT_QUOTE_BORDER_COLOR)
        _set_quote_border(paragraph, color)
        try:
            paragraph_format.left_indent = Inches(0.2)
        except Exception:
            pass


def _add_markdown_content(document, md_text: str, opts: Dict[str, Any]) -> None:
    try:
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        pass
    
    # 处理连续标题和引用块的情况，记录上下文
    is_after_heading = False
    lines = md_text.splitlines()
    i = 0
    
    while i < len(lines):
        line = lines[i]
        paragraph = None
        is_quote = False
        
        # 改进标题行识别逻辑，确保"###"开头的行被正确识别为标题，不会被误认为列表项
        if (line.startswith("### ") or 
            line.startswith("## ") or 
            line.startswith("# ")):
            
            # 确定标题级别
            level = 1 if line.startswith("# ") else (2 if line.startswith("## ") else 3)
            title_text = line[level + 1:]  # 去掉 # 符号和空格
            
            # 避免将标题误识别为列表项 - 显式使用heading样式
            paragraph = document.add_heading(title_text, level=level)
            
            # 特殊处理说话人标题
            if level == 3:
                try:
                    for run in paragraph.runs:
                        run.font.name = "宋体"
                        if hasattr(run._element, "rPr") and hasattr(run._element.rPr, "rFonts"):
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                except Exception:
                    pass
            
            is_after_heading = True
            i += 1
            
            # 如果下一行是空行且再下一行是引用块，则跳过空行，实现无间隔处理
            if i < len(lines) - 1 and not lines[i].strip() and i + 1 < len(lines) and lines[i+1].startswith(">"):
                i += 1  # 跳过空行
            
            continue
                
        # 处理引用块
        elif line.startswith(">"):
            is_quote = True
            text = line.lstrip(">").lstrip()
            
            try:
                # 创建一个引用段落，使用Quote样式但不应用斜体
                paragraph = document.add_paragraph(text, style="Quote")
                # 遍历段落中的所有文本运行(runs)，确保没有应用斜体
                for run in paragraph.runs:
                    run.italic = False
                    
                # 如果是紧跟标题的引用块，减少段前间距
                if is_after_heading and paragraph.paragraph_format:
                    paragraph.paragraph_format.space_before = 0
            except Exception:
                paragraph = document.add_paragraph(text)
                for run in paragraph.runs:
                    run.italic = False
                    
            is_after_heading = False
        
        # 处理列表项（自定义格式，不使用pandoc默认样式）
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:].strip()
            paragraph = document.add_paragraph()
            # 添加自定义项目符号而不是使用默认列表样式
            run = paragraph.add_run("• ")
            paragraph.add_run(text)
            paragraph.paragraph_format.left_indent = document.styles['Normal'].paragraph_format.left_indent
            is_after_heading = False
        
        elif line.strip().startswith("1. ") or re.match(r"^\d+\.\s+", line.strip()):
            # 处理有序列表
            match = re.match(r"^(\d+)\.\s+(.*)", line.strip())
            if match:
                num, text = match.groups()
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"{num}. ")
                paragraph.add_run(text)
                paragraph.paragraph_format.left_indent = document.styles['Normal'].paragraph_format.left_indent
            is_after_heading = False
        
        elif line.strip():
            paragraph = document.add_paragraph(line)
            is_after_heading = False
        
        else:
            paragraph = document.add_paragraph("")
            # 保持标题后的空行状态，不重置标记
            # is_after_heading = False

        if paragraph is not None:
            _apply_paragraph_format(paragraph, opts, is_quote=is_quote)
        
        i += 1


def _ensure_quote_style_without_italic(document) -> None:
    """确保Quote样式存在且不包含斜体设置"""
    try:
        from docx.shared import Inches  # type: ignore
        # 尝试获取Quote样式，如果不存在则创建
        try:
            quote_style = document.styles["Quote"]
        except KeyError:
            quote_style = document.styles.add_style("Quote", 1)  # 1表示段落样式
            
        # 确保样式不包含斜体设置
        quote_style.font.italic = False
        
        # 设置引用段落的段前段后间距较小
        if hasattr(quote_style, "paragraph_format"):
            quote_style.paragraph_format.space_before = 0
            quote_style.paragraph_format.space_after = 0
            quote_style.paragraph_format.left_indent = Inches(0.2)
        
        return quote_style
    except Exception:
        # 如果无法修改样式，则静默失败，继续使用默认处理
        pass
    return None


def _docx_bytes_from_markdown(md_text: str, title: Optional[str], options: Optional[DocxStyleOptions]) -> bytes:
    try:
        from docx import Document as DocxDocument  # type: ignore
        from docx.enum.style import WD_STYLE_TYPE  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"docx 依赖缺失: {exc}")

    document = DocxDocument()
    opts = options.dict(exclude_none=True) if options else {}
    _apply_normal_style(document, opts)
    _apply_heading_styles(document, opts)
    
    # 确保Quote样式正确设置，不使用斜体
    _ensure_quote_style_without_italic(document)
    
    # 修改列表样式，防止使用默认项目符号
    try:
        # 获取列表样式并修改其属性
        list_styles = [s for s in document.styles if s.type == WD_STYLE_TYPE.LIST]
        for list_style in list_styles:
            try:
                # 尝试修改列表格式，简化项目符号
                if hasattr(list_style, "_element"):
                    # 可以在这里修改列表样式的详细属性
                    pass
            except Exception:
                pass
    except Exception:
        # 如果无法修改列表样式，静默失败
        pass
    
    # 检查markdown文本是否已经包含标题，避免重复
    lines = md_text.split('\n')
    has_title_in_md = False
    
    if title and lines and lines[0].startswith('# '):
        # 如果Markdown已有标题行且与请求标题匹配，移除markdown中的标题以避免重复
        md_title = lines[0][2:].strip()
        if md_title == title:
            md_text = '\n'.join(lines[1:]).strip()
            has_title_in_md = True
    
    # 添加标题，确保使用等线字体
    if title:
        heading = document.add_heading(title, level=0)
        # 为标题单独设置等线字体
        try:
            from docx.oxml.ns import qn
            for run in heading.runs:
                run.font.name = "等线"
                if hasattr(run._element, "rPr") and hasattr(run._element.rPr, "rFonts"):
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        except Exception:
            pass
        _apply_paragraph_format(heading, opts)

    # 使用改进的方法添加内容，避免pandoc默认的项目符号格式和间距问题
    _add_markdown_content(document, md_text, opts)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _which(binary: str) -> str:
    return shutil.which(binary) or ""


def _load_markdown_css(for_pdf: bool = False) -> str:
    backend_css_paths = [
        Path(__file__).resolve().parent.parent.parent / "assets" / MARKDOWN_CSS_NAME,
        Path(os.getcwd()) / "backend" / "assets" / MARKDOWN_CSS_NAME,
    ]
    css_content = ""
    for candidate in backend_css_paths:
        if candidate.exists():
            try:
                css_content = candidate.read_text(encoding="utf-8")
                break
            except Exception:
                continue

    if for_pdf:
        base = css_content or ""
        overrides = (
            "body{background:#ffffff!important;}"
            # 修改PDF的字体为宋体
            ".markdown{background:#ffffff!important;color:#111827;font-family:宋体,SimSun,'Noto Sans CJK SC',serif;}"
            ".markdown blockquote{background:transparent!important;border-left:4px solid #e5e7eb!important;padding-left:12px;margin:12px 0;color:#6b7280;}"
            ".markdown h1{font-size:28pt!important;font-weight:700;margin:0 0 12pt 0;}"
        )
        return base + "\n" + overrides

    frontend_css_candidates = [
        Path(os.getcwd()) / "frontend" / "src" / "styles.css",
        Path(__file__).resolve().parents[4] / "frontend" / "src" / "styles.css",
    ]
    for candidate in frontend_css_candidates:
        if candidate.exists():
            try:
                extra_css = candidate.read_text(encoding="utf-8")
                return css_content + "\n" + extra_css
            except Exception:
                continue

    return css_content or (
        # 修改默认CSS中的字体为宋体
        ".markdown{font-family:宋体,SimSun,-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;color:#111;line-height:1.6;}\n"
        ".markdown h1,.markdown h2,.markdown h3{margin-top:1.2em;}\n"
        ".markdown p{line-height:1.6;}\n"
        "@page{size:A4;margin:20mm;}\n"
        "body{margin:0;}\n"
    )


def _markdown_to_html(md_text: str, title: Optional[str], for_pdf: bool = False) -> str:
    # 添加sane_lists扩展，确保更严格的列表识别规则
    html = mdlib.markdown(md_text, extensions=["extra", "sane_lists", "toc", "tables"])
    css = _load_markdown_css(for_pdf=for_pdf)
    title_html = f"<title>{title}</title>" if title else ""
    
    # 添加自定义CSS规则，确保###开头的内容正确显示为标题
    additional_css = """
    h3 { display: block; font-size: 1.17em; margin-block-start: 1em; margin-block-end: 1em; font-weight: bold; }
    """
    
    return (
        "<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'>"
        f"{title_html}<style>{css}{additional_css}</style></head><body><div class='markdown'>{html}</div></body></html>"
    )


def _html_to_pdf_via_chrome(html: str) -> bytes:
    chrome = (
        os.getenv("CHROME_BIN")
        or _which("google-chrome-stable")
        or _which("google-chrome")
        or _which("chromium-browser")
        or _which("chromium")
    )
    
    # 记录更详细的Chrome路径信息，帮助诊断
    print(f"尝试使用Chrome导出PDF, 路径: {chrome}")
    
    if not chrome:
        raise HTTPException(status_code=500, detail="未找到 Chrome/Chromium，可设置 CHROME_BIN 环境变量或安装浏览器。")

    with tempfile.TemporaryDirectory(prefix="wordline_pdf_") as temp_dir:
        html_path = Path(temp_dir) / "input.html"
        pdf_path = Path(temp_dir) / "output.pdf"
        html_path.write_text(html, encoding="utf-8")
        url = f"file://{html_path}"
        
        # 记录临时文件路径
        print(f"临时HTML文件: {html_path}, 输出PDF路径: {pdf_path}")
        
        try:
            # 构建命令并记录
            cmd = [
                chrome,
                "--headless=new",
                "--disable-gpu",
                "--no-sandbox",
                f"--print-to-pdf={pdf_path}",
                url,
            ]
            print(f"执行Chrome命令: {' '.join(cmd)}")
            
            # 收集标准输出和错误输出以便调试
            result = subprocess.run(
                cmd,
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )
            
            # 记录命令的输出
            if result.stdout:
                print(f"Chrome stdout: {result.stdout.decode('utf-8', errors='replace')}")
            if result.stderr:
                print(f"Chrome stderr: {result.stderr.decode('utf-8', errors='replace')}")
                
        except subprocess.TimeoutExpired as exc:
            print(f"Chrome命令超时: {exc}")
            raise HTTPException(status_code=500, detail=f"Chrome headless 执行超时: {exc}")
        except subprocess.CalledProcessError as exc:
            stdout = exc.stdout.decode("utf-8", errors="replace") if exc.stdout else ""
            stderr = exc.stderr.decode("utf-8", errors="replace") if exc.stderr else ""
            print(f"Chrome命令失败: {exc}. stdout: {stdout}, stderr: {stderr}")
            raise HTTPException(status_code=500, detail=f"Chrome headless 执行失败: {stderr or str(exc)}")
        except Exception as exc:  # noqa: BLE001
            print(f"Chrome导出PDF过程中出现未知错误: {exc}")
            raise HTTPException(status_code=500, detail=f"Chrome headless 失败: {exc}")
            
        if not pdf_path.exists():
            print(f"Chrome未能生成PDF文件: {pdf_path}")
            raise HTTPException(status_code=500, detail="Chrome 未生成 PDF 文件")
            
        try:
            pdf_data = pdf_path.read_bytes()
            print(f"成功读取PDF文件，大小: {len(pdf_data)} 字节")
            return pdf_data
        except Exception as exc:
            print(f"读取生成的PDF文件失败: {exc}")
            raise HTTPException(status_code=500, detail=f"读取生成的PDF文件失败: {exc}")


def _html_to_pdf_via_wkhtmltopdf(html: str) -> bytes:
    wkhtml = os.getenv("WKHTMLTOPDF_BIN") or _which("wkhtmltopdf")
    
    # 记录更详细的wkhtmltopdf路径信息
    print(f"尝试使用wkhtmltopdf导出PDF, 路径: {wkhtml}")
    
    if not wkhtml:
        raise HTTPException(status_code=500, detail="未找到 wkhtmltopdf，可设置 WKHTMLTOPDF_BIN 环境变量或安装系统包。")

    with tempfile.TemporaryDirectory(prefix="wordline_pdf_") as temp_dir:
        html_path = Path(temp_dir) / "input.html"
        pdf_path = Path(temp_dir) / "output.pdf"
        html_path.write_text(html, encoding="utf-8")
        
        # 记录临时文件路径
        print(f"临时HTML文件: {html_path}, 输出PDF路径: {pdf_path}")
        
        try:
            # 构建命令并记录
            cmd = [
                wkhtml,
                "-s",
                "A4",
                "--margin-top", "20mm",
                "--margin-bottom", "20mm",
                "--margin-left", "20mm",
                "--margin-right", "20mm",
                "--disable-smart-shrinking",
                str(html_path),
                str(pdf_path),
            ]
            print(f"执行wkhtmltopdf命令: {' '.join(cmd)}")
            
            # 收集标准输出和错误输出以便调试
            result = subprocess.run(
                cmd,
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )
            
            # 记录命令的输出
            if result.stdout:
                print(f"wkhtmltopdf stdout: {result.stdout.decode('utf-8', errors='replace')}")
            if result.stderr:
                print(f"wkhtmltopdf stderr: {result.stderr.decode('utf-8', errors='replace')}")
                
        except subprocess.TimeoutExpired as exc:
            print(f"wkhtmltopdf命令超时: {exc}")
            raise HTTPException(status_code=500, detail=f"wkhtmltopdf 执行超时: {exc}")
        except subprocess.CalledProcessError as exc:
            stdout = exc.stdout.decode("utf-8", errors="replace") if exc.stdout else ""
            stderr = exc.stderr.decode("utf-8", errors="replace") if exc.stderr else ""
            print(f"wkhtmltopdf命令失败: {exc}. stdout: {stdout}, stderr: {stderr}")
            raise HTTPException(status_code=500, detail=f"wkhtmltopdf 执行失败: {stderr or str(exc)}")
        except Exception as exc:  # noqa: BLE001
            print(f"wkhtmltopdf导出PDF过程中出现未知错误: {exc}")
            raise HTTPException(status_code=500, detail=f"wkhtmltopdf 失败: {exc}")
            
        if not pdf_path.exists():
            print(f"wkhtmltopdf未能生成PDF文件: {pdf_path}")
            raise HTTPException(status_code=500, detail="wkhtmltopdf 未生成 PDF 文件")
            
        try:
            pdf_data = pdf_path.read_bytes()
            print(f"成功读取PDF文件，大小: {len(pdf_data)} 字节")
            return pdf_data
        except Exception as exc:
            print(f"读取生成的PDF文件失败: {exc}")
            raise HTTPException(status_code=500, detail=f"读取生成的PDF文件失败: {exc}")


def _detect_engines() -> Dict[str, Dict[str, Any]]:
    chrome_paths = [
        os.getenv("CHROME_BIN"),
        _which("google-chrome-stable"),
        _which("google-chrome"),
        _which("chromium-browser"),
        _which("chromium"),
    ]
    chrome_found = any(path for path in chrome_paths)
    chrome_path = next((path for path in chrome_paths if path), "")
    
    wkhtml_path = os.getenv("WKHTMLTOPDF_BIN") or _which("wkhtmltopdf")
    
    print(f"检测到的PDF引擎: Chrome={chrome_found}({chrome_path}), wkhtmltopdf={bool(wkhtml_path)}({wkhtml_path})")
    
    return {
        "chrome": {"found": chrome_found, "bin": chrome_path},
        "wkhtmltopdf": {"found": bool(wkhtml_path), "bin": wkhtml_path},
    }


def _respond_markdown(markdown_text: str) -> PlainTextResponse:
    return PlainTextResponse(content=markdown_text, media_type="text/markdown")


def _respond_docx(markdown_text: str, title: Optional[str], options: Optional[DocxStyleOptions]) -> StreamingResponse:
    data = _docx_bytes_from_markdown(markdown_text, title, options)
    buffer = io.BytesIO(data)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="export.docx"'},
    )


def _html_to_pdf_via_reportlab(markdown_text: str, title: Optional[str]) -> bytes:
    """使用reportlab生成PDF（纯Python实现，无需外部工具）"""
    print("使用reportlab生成PDF（备用方式）")
    
    if not HAS_REPORTLAB:
        raise HTTPException(status_code=500, detail="reportlab库未安装，无法使用备用PDF生成")
    
    buffer = io.BytesIO()
    
    # 注册中文字体
    font_name = _register_cjk_fonts() or "Helvetica"  # 如果没有中文字体，使用默认的Helvetica
    
    # 创建文档
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        leftMargin=20*mm, 
        rightMargin=20*mm,
        topMargin=20*mm,
        bottomMargin=20*mm
    )
    
    # 创建样式
    styles = getSampleStyleSheet()
    
    # 添加中文支持的样式
    if font_name != "Helvetica":
        styles.add(ParagraphStyle(
            name='ChineseBody',
            fontName=font_name,
            fontSize=12,
            leading=14,
            spaceAfter=12
        ))
        styles.add(ParagraphStyle(
            name='ChineseTitle',
            fontName=font_name,
            fontSize=18,
            leading=22,
            spaceAfter=24,
            alignment=1  # 居中
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading',
            fontName=font_name,
            fontSize=14,
            leading=16,
            spaceAfter=12
        ))
        styles.add(ParagraphStyle(
            name='ChineseBlockquote',
            fontName=font_name,
            fontSize=12,
            leading=14,
            leftIndent=20,
            rightIndent=20,
            spaceAfter=12
        ))
    
    # 转换markdown为简单格式
    content = []
    
    # 添加标题
    if title:
        style_name = 'ChineseTitle' if font_name != "Helvetica" else 'Title'
        content.append(Paragraph(title, styles[style_name]))
        content.append(Spacer(1, 12))
    
    # 将Markdown拆分为段落，简单处理标题和引用块
    for line in markdown_text.split('\n'):
        if not line.strip():
            continue
            
        # 处理标题
        if line.startswith('### '):
            style_name = 'ChineseHeading' if font_name != "Helvetica" else 'Heading3'
            content.append(Paragraph(line[4:], styles[style_name]))
        elif line.startswith('## '):
            style_name = 'ChineseHeading' if font_name != "Helvetica" else 'Heading2'
            content.append(Paragraph(line[3:], styles[style_name]))
        elif line.startswith('# '):
            style_name = 'ChineseHeading' if font_name != "Helvetica" else 'Heading1'
            content.append(Paragraph(line[2:], styles[style_name]))
        # 处理引用块
        elif line.startswith('> '):
            style_name = 'ChineseBlockquote' if font_name != "Helvetica" else 'BodyText'
            content.append(Paragraph(line[2:], styles[style_name]))
        # 普通段落
        else:
            style_name = 'ChineseBody' if font_name != "Helvetica" else 'BodyText'
            content.append(Paragraph(line, styles[style_name]))
    
    # 构建PDF
    try:
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"生成PDF失败: {e}")
        raise HTTPException(status_code=500, detail=f"生成PDF失败: {e}")


def _respond_pdf(markdown_text: str, title: Optional[str]) -> StreamingResponse:
    print("开始PDF导出流程...")
    html = _markdown_to_html(markdown_text, title, for_pdf=True)
    print(f"已生成HTML，长度: {len(html)} 字符")
    
    # 先检查有哪些可用的引擎
    engines = _detect_engines()
    print(f"可用的PDF引擎: {engines}")
    
    errors: list[str] = []
    pdf_bytes: Optional[bytes] = None
    
    # 尝试使用Chrome
    if engines["chrome"]["found"]:
        try:
            print("尝试使用Chrome生成PDF...")
            pdf_bytes = _html_to_pdf_via_chrome(html)
            print(f"Chrome成功生成PDF，大小: {len(pdf_bytes) if pdf_bytes else 0} 字节")
        except HTTPException as exc:
            error_msg = str(exc.detail)
            print(f"使用Chrome生成PDF失败: {error_msg}")
            errors.append(error_msg)
        except Exception as exc:  # noqa: BLE001
            error_msg = f"chrome: {exc}"
            print(f"使用Chrome生成PDF时发生异常: {error_msg}")
            errors.append(error_msg)

    # 如果Chrome失败，尝试wkhtmltopdf
    if pdf_bytes is None and engines["wkhtmltopdf"]["found"]:
        try:
            print("尝试使用wkhtmltopdf生成PDF...")
            pdf_bytes = _html_to_pdf_via_wkhtmltopdf(html)
            print(f"wkhtmltopdf成功生成PDF，大小: {len(pdf_bytes) if pdf_bytes else 0} 字节")
        except HTTPException as exc:
            error_msg = str(exc.detail)
            print(f"使用wkhtmltopdf生成PDF失败: {error_msg}")
            errors.append(error_msg)
        except Exception as exc:  # noqa: BLE001
            error_msg = f"wkhtmltopdf: {exc}"
            print(f"使用wkhtmltopdf生成PDF时发生异常: {error_msg}")
            errors.append(error_msg)

    # 尝试使用reportlab作为备用方案
    if pdf_bytes is None and HAS_REPORTLAB:
        try:
            print("尝试使用reportlab(纯Python)生成PDF...")
            pdf_bytes = _html_to_pdf_via_reportlab(markdown_text, title)
            print(f"reportlab成功生成PDF，大小: {len(pdf_bytes) if pdf_bytes else 0} 字节")
        except HTTPException as exc:
            error_msg = str(exc.detail)
            print(f"使用reportlab生成PDF失败: {error_msg}")
            errors.append(error_msg)
        except Exception as exc:
            error_msg = f"reportlab: {exc}"
            print(f"使用reportlab生成PDF时发生异常: {error_msg}")
            errors.append(error_msg)

    # 所有方法都失败
    if pdf_bytes is None:
        detail = "无法生成PDF"
        # 检查是否没有可用的引擎
        if not engines["chrome"]["found"] and not engines["wkhtmltopdf"]["found"] and not HAS_REPORTLAB:
            detail += "：未检测到可用的PDF导出引擎。请安装Chrome/Chromium或wkhtmltopdf，或者安装reportlab库(pip install reportlab)"
        elif errors:
            detail += "：" + "; ".join(errors)
            
        print(f"所有PDF生成方法均失败: {detail}")
        raise HTTPException(status_code=500, detail=detail)

    buffer = io.BytesIO(pdf_bytes)
    print("PDF导出成功，准备返回响应")
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="export.pdf"'},
    )


@router.get("/engines")
def engines():
    engines = _detect_engines()
    # 添加reportlab作为备用选项
    engines["reportlab"] = {"found": HAS_REPORTLAB, "bin": "内置Python库"}
    return engines


@router.post("")
def export(req: ExportRequest, fmt: str = Query(pattern=r"^(md|docx|pdf)$")):
    fmt = fmt.lower()
    
    # 过滤内部处理消息、空内容块和日期时间格式的块
    filtered_blocks = []
    from src.services.markdowner import is_internal_message, is_likely_timestamp
    
    for block in req.blocks:
        content = block.get("content", "").strip()
        speaker = block.get("speaker", "").strip()
        
        # 排除内部处理消息和空内容
        if content and not is_internal_message(content):
            # 检查内容中是否含有常见的内部处理消息模式
            internal_patterns = [
                r"^当前段落.*[:：]",
                r"^\[.*?\]$",
                r"^前文：.*?后文：",
            ]
            
            is_internal = any(re.search(pattern, content) for pattern in internal_patterns)
            if is_internal:
                continue
            
            # 检查说话人是否为日期时间格式    
            if speaker and is_likely_timestamp(speaker):
                timestamp = block.get("timestamp", "")
                # 如果说话人是时间格式，则移到timestamp字段
                if not timestamp:
                    block["timestamp"] = speaker
                    block["speaker"] = ""
                else:
                    # 如果已有timestamp，且speaker是日期格式，直接清空speaker
                    block["speaker"] = ""
                    
            # 额外检查是否整个标题是日期格式
            if block["speaker"] == "" and block.get("timestamp", "") and is_likely_timestamp(block["timestamp"]):
                block["timestamp"] = ""
                    
            filtered_blocks.append(block)
    
    # 使用过滤后的块生成markdown文本
    export_request = ExportRequest(blocks=filtered_blocks, title=req.title, docx_options=req.docx_options)
    markdown_text = blocks_to_markdown(export_request.blocks, title=export_request.title)
    
    handlers = {
        "md": lambda: _respond_markdown(markdown_text),
        "docx": lambda: _respond_docx(markdown_text, export_request.title, export_request.docx_options),
        "pdf": lambda: _respond_pdf(markdown_text, export_request.title),
    }

    try:
        return handlers[fmt]()
    except KeyError as exc:
        raise HTTPException(status_code=400, detail="不支持的导出格式") from exc
